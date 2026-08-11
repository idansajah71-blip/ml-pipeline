"""Export Service — Multi-format data export from scraped/processed data.
Supports CSV, Excel, JSON, Word (.docx), HTML, Parquet, SQL, XML."""
import io
import json
import os
from typing import Optional, List, Dict, Any
from datetime import datetime

import pandas as pd


class ExportService:

    EXPORT_DIR = "ml_artifacts/exports"

    def __init__(self):
        os.makedirs(self.EXPORT_DIR, exist_ok=True)

    def export_csv(self, df: pd.DataFrame, filename: str = None) -> dict:
        if filename is None:
            filename = f"export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
        filepath = os.path.join(self.EXPORT_DIR, filename)
        df.to_csv(filepath, index=False)
        csv_string = df.to_csv(index=False)
        return {
            "format": "csv",
            "filepath": filepath,
            "filename": filename,
            "size_bytes": os.path.getsize(filepath),
            "row_count": len(df),
            "column_count": len(df.columns),
            "content": csv_string[:10000] if len(csv_string) <= 10000 else csv_string[:10000] + "...",
        }

    def export_excel(self, df: pd.DataFrame, filename: str = None, sheet_name: str = "Sheet1") -> dict:
        if filename is None:
            filename = f"export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        filepath = os.path.join(self.EXPORT_DIR, filename)
        df.to_excel(filepath, index=False, sheet_name=sheet_name, engine="openpyxl")
        return {
            "format": "excel",
            "filepath": filepath,
            "filename": filename,
            "size_bytes": os.path.getsize(filepath),
            "row_count": len(df),
            "column_count": len(df.columns),
            "sheet_name": sheet_name,
        }

    def export_json(self, df: pd.DataFrame, filename: str = None, orient: str = "records") -> dict:
        if filename is None:
            filename = f"export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        filepath = os.path.join(self.EXPORT_DIR, filename)
        data = df.where(df.notna(), None).to_dict(orient=orient)
        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, default=str, ensure_ascii=False)
        return {
            "format": "json",
            "filepath": filepath,
            "filename": filename,
            "size_bytes": os.path.getsize(filepath),
            "row_count": len(df),
            "column_count": len(df.columns),
            "orient": orient,
        }

    def export_word(self, df: pd.DataFrame, filename: str = None) -> dict:
        """Export DataFrame to a Word (.docx) file with a formatted table."""
        from docx import Document
        from docx.shared import Inches

        if filename is None:
            filename = f"export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.EXPORT_DIR, filename)

        doc = Document()
        doc.add_heading("Scraped Data Export", 0)

        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)

        for _, row in df.head(5000).iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(df.columns):
                val = row[col]
                cells[i].text = "" if pd.isna(val) else str(val)

        doc.add_paragraph(f"\nTotal rows: {len(df)}")
        doc.add_paragraph(f"Columns: {len(df.columns)}")
        doc.add_paragraph(f"Exported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        doc.save(filepath)
        return {
            "format": "word",
            "filepath": filepath,
            "filename": filename,
            "size_bytes": os.path.getsize(filepath),
            "row_count": min(len(df), 5000),
            "column_count": len(df.columns),
        }

    def export_html(self, df: pd.DataFrame, filename: str = None) -> dict:
        """Export DataFrame to an HTML table file."""
        if filename is None:
            filename = f"export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
        filepath = os.path.join(self.EXPORT_DIR, filename)

        safe = df.where(df.notna(), None)
        html = safe.to_html(index=False, border=1, escape=False)
        full = f"<!DOCTYPE html><html><head><meta charset='utf-8'><title>Scraped Data</title></head><body>{html}</body></html>"
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(full)

        return {
            "format": "html",
            "filepath": filepath,
            "filename": filename,
            "size_bytes": os.path.getsize(filepath),
            "row_count": len(df),
            "column_count": len(df.columns),
        }

    def export_parquet(self, df: pd.DataFrame, filename: str = None) -> dict:
        if filename is None:
            filename = f"export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.parquet"
        filepath = os.path.join(self.EXPORT_DIR, filename)
        df.to_parquet(filepath, index=False, engine="pyarrow")
        return {
            "format": "parquet",
            "filepath": filepath,
            "filename": filename,
            "size_bytes": os.path.getsize(filepath),
            "row_count": len(df),
            "column_count": len(df.columns),
        }

    def export_xml(self, df: pd.DataFrame, filename: str = None, root_tag: str = "data", row_tag: str = "record") -> dict:
        if filename is None:
            filename = f"export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xml"
        filepath = os.path.join(self.EXPORT_DIR, filename)
        lines = [f'<?xml version="1.0" encoding="UTF-8"?>', f'<{root_tag}>']
        for _, row in df.iterrows():
            lines.append(f'  <{row_tag}>')
            for col in df.columns:
                val = row[col]
                if pd.isna(val):
                    val = ""
                else:
                    val = str(val).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                safe_col = col.replace(" ", "_").replace("-", "_")
                lines.append(f'    <{safe_col}>{val}</{safe_col}>')
            lines.append(f'  </{row_tag}>')
        lines.append(f'</{root_tag}>')
        xml_content = "\n".join(lines)
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(xml_content)
        return {
            "format": "xml",
            "filepath": filepath,
            "filename": filename,
            "size_bytes": os.path.getsize(filepath),
            "row_count": len(df),
            "column_count": len(df.columns),
        }

    def export_sql(self, df: pd.DataFrame, table_name: str = "scraped_data", filename: str = None) -> dict:
        if filename is None:
            filename = f"export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.sql"
        filepath = os.path.join(self.EXPORT_DIR, filename)
        lines = []
        col_defs = []
        for col in df.columns:
            safe_col = col.replace(" ", "_").replace("-", "_")
            sample = df[col].dropna().head(20)
            if pd.api.types.is_numeric_dtype(df[col]):
                if pd.api.types.is_integer_dtype(df[col]):
                    col_defs.append(f"  {safe_col} INTEGER")
                else:
                    col_defs.append(f"  {safe_col} NUMERIC")
            else:
                max_len = sample.astype(str).str.len().max() if len(sample) > 0 else 255
                col_len = max(50, min(int(max_len * 1.5), 4000))
                col_defs.append(f"  {safe_col} VARCHAR({col_len})")

        lines.append(f"CREATE TABLE IF NOT EXISTS {table_name} (")
        lines.append(",\n".join(col_defs))
        lines.append(");\n")

        for _, row in df.head(1000).iterrows():
            values = []
            for col in df.columns:
                val = row[col]
                if pd.isna(val):
                    values.append("NULL")
                elif isinstance(val, (int, float)):
                    values.append(str(val))
                else:
                    escaped = str(val).replace("'", "''")
                    values.append(f"'{escaped}'")
            safe_cols = [c.replace(" ", "_").replace("-", "_") for c in df.columns]
            lines.append(
                f"INSERT INTO {table_name} ({', '.join(safe_cols)}) VALUES ({', '.join(values)});"
            )

        sql_content = "\n".join(lines)
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(sql_content)
        return {
            "format": "sql",
            "filepath": filepath,
            "filename": filename,
            "size_bytes": os.path.getsize(filepath),
            "row_count": min(len(df), 1000),
            "table_name": table_name,
        }

    def export_pdf(self, df: pd.DataFrame, filename: str = None, title: str = "Scraped Data Export") -> dict:
        """Export DataFrame to a paginated PDF report using ReportLab."""
        try:
            from reportlab.lib.pagesizes import A4, landscape
            from reportlab.lib import colors
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.platypus import (
                SimpleDocTemplate, Table, TableStyle, Paragraph,
                Spacer, HRFlowable,
            )
            from reportlab.lib.enums import TA_CENTER
        except ImportError:
            raise ImportError(
                "reportlab is required for PDF export. "
                "Install it with: pip install reportlab==4.2.2"
            )

        if filename is None:
            filename = f"export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        filepath = os.path.join(self.EXPORT_DIR, filename)

        # Use landscape for wide tables
        ncols = len(df.columns)
        page_size = landscape(A4) if ncols > 6 else A4
        doc = SimpleDocTemplate(
            filepath,
            pagesize=page_size,
            leftMargin=1.5 * cm,
            rightMargin=1.5 * cm,
            topMargin=1.5 * cm,
            bottomMargin=1.5 * cm,
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "ReportTitle",
            parent=styles["Heading1"],
            alignment=TA_CENTER,
            spaceAfter=0.4 * cm,
            fontSize=16,
        )
        meta_style = ParagraphStyle(
            "Meta",
            parent=styles["Normal"],
            alignment=TA_CENTER,
            fontSize=9,
            textColor=colors.HexColor("#6b7280"),
            spaceAfter=0.6 * cm,
        )

        story = []
        story.append(Paragraph(title, title_style))
        story.append(Paragraph(
            f"{len(df):,} baris · {len(df.columns)} kolom · "
            f"Diekspor {datetime.now().strftime('%d %b %Y %H:%M')}",
            meta_style,
        ))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#e5e7eb")))
        story.append(Spacer(1, 0.4 * cm))

        # Cap rows to keep PDF reasonable
        max_rows = 2000
        display_df = df.head(max_rows)

        # Build table data: header + rows
        col_names = [str(c)[:25] for c in display_df.columns]  # truncate long headers
        data = [col_names]
        for _, row in display_df.iterrows():
            data.append([
                ("" if (v is None or (isinstance(v, float) and v != v)) else str(v)[:60])
                for v in row
            ])

        # Distribute column widths evenly within page
        page_w, _ = page_size
        usable_w = page_w - 3 * cm
        col_w = min(usable_w / max(ncols, 1), 6 * cm)
        col_widths = [col_w] * ncols

        tbl = Table(data, colWidths=col_widths, repeatRows=1)
        tbl.setStyle(TableStyle([
            # Header
            ("BACKGROUND",   (0, 0), (-1, 0),  colors.HexColor("#4f46e5")),
            ("TEXTCOLOR",    (0, 0), (-1, 0),  colors.white),
            ("FONTNAME",     (0, 0), (-1, 0),  "Helvetica-Bold"),
            ("FONTSIZE",     (0, 0), (-1, 0),  8),
            ("BOTTOMPADDING",(0, 0), (-1, 0),  6),
            ("TOPPADDING",   (0, 0), (-1, 0),  6),
            # Body
            ("FONTSIZE",     (0, 1), (-1, -1), 7),
            ("TOPPADDING",   (0, 1), (-1, -1), 4),
            ("BOTTOMPADDING",(0, 1), (-1, -1), 4),
            # Alternating row colours
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f9fafb")]),
            # Grid
            ("GRID",         (0, 0), (-1, -1), 0.3, colors.HexColor("#d1d5db")),
            ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
            ("WORDWRAP",     (0, 0), (-1, -1), True),
        ]))
        story.append(tbl)

        if len(df) > max_rows:
            story.append(Spacer(1, 0.4 * cm))
            story.append(Paragraph(
                f"⚠ Hanya menampilkan {max_rows:,} dari {len(df):,} baris.",
                ParagraphStyle("warn", parent=styles["Normal"], fontSize=8,
                               textColor=colors.HexColor("#f59e0b")),
            ))

        doc.build(story)
        return {
            "format": "pdf",
            "filepath": filepath,
            "filename": filename,
            "size_bytes": os.path.getsize(filepath),
            "row_count": min(len(df), max_rows),
            "column_count": len(df.columns),
            "truncated": len(df) > max_rows,
        }

    def export_to_buffer(self, df: pd.DataFrame, fmt: str, **kwargs) -> tuple[bytes, str]:
        """
        Export DataFrame to an in-memory buffer for streaming HTTP responses.
        Returns (bytes_content, mime_type).
        """
        buf = io.BytesIO()
        mime = "application/octet-stream"

        if fmt == "csv":
            content = df.to_csv(index=False).encode("utf-8")
            return content, "text/csv; charset=utf-8"

        if fmt == "json":
            content = df.where(df.notna(), None).to_json(
                orient="records", indent=2, force_ascii=False
            ).encode("utf-8")
            return content, "application/json; charset=utf-8"

        if fmt == "excel":
            df.to_excel(buf, index=False, engine="openpyxl")
            return buf.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

        if fmt == "parquet":
            df.to_parquet(buf, index=False, engine="pyarrow")
            return buf.getvalue(), "application/octet-stream"

        if fmt == "xml":
            root_tag = kwargs.get("root_tag", "data")
            row_tag  = kwargs.get("row_tag", "record")
            lines = [f'<?xml version="1.0" encoding="UTF-8"?>', f"<{root_tag}>"]
            for _, row in df.iterrows():
                lines.append(f"  <{row_tag}>")
                for col in df.columns:
                    v = row[col]
                    v = "" if (v is None or (isinstance(v, float) and v != v)) else str(v)
                    v = v.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    safe_col = col.replace(" ", "_").replace("-", "_")
                    lines.append(f"    <{safe_col}>{v}</{safe_col}>")
                lines.append(f"  </{row_tag}>")
            lines.append(f"</{root_tag}>")
            return "\n".join(lines).encode("utf-8"), "application/xml; charset=utf-8"

        if fmt in ("word", "docx"):
            from docx import Document
            doc = Document()
            doc.add_heading("Scraped Data Export", 0)
            tbl = doc.add_table(rows=1, cols=len(df.columns))
            tbl.style = "Table Grid"
            for i, col in enumerate(df.columns):
                tbl.rows[0].cells[i].text = str(col)
            for _, row in df.head(5000).iterrows():
                cells = tbl.add_row().cells
                for i, col in enumerate(df.columns):
                    v = row[col]
                    cells[i].text = "" if (v is None or str(v) in ("nan", "None")) else str(v)
            doc.add_paragraph(f"\nTotal: {len(df)} baris · {len(df.columns)} kolom")
            doc.add_paragraph(f"Diekspor: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.save(buf)
            return buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        if fmt == "pdf":
            # Write to temp file (reportlab requires a real path), then read back
            import tempfile
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp_path = tmp.name
            try:
                self.export_pdf(df, filename=os.path.basename(tmp_path), title="Scraped Data Export")
                real_path = os.path.join(self.EXPORT_DIR, os.path.basename(tmp_path))
                with open(real_path, "rb") as f:
                    content = f.read()
                os.unlink(real_path)
            finally:
                if os.path.exists(tmp_path):
                    os.unlink(tmp_path)
            return content, "application/pdf"

        if fmt == "sql":
            lines = []
            col_defs = []
            for col in df.columns:
                safe_col = col.replace(" ", "_").replace("-", "_")
                if pd.api.types.is_integer_dtype(df[col]):
                    col_defs.append(f"  {safe_col} INTEGER")
                elif pd.api.types.is_float_dtype(df[col]):
                    col_defs.append(f"  {safe_col} NUMERIC")
                else:
                    col_defs.append(f"  {safe_col} TEXT")
            lines += [
                "CREATE TABLE IF NOT EXISTS scraped_data (", ",\n".join(col_defs), ");\n"
            ]
            for _, row in df.head(1000).iterrows():
                vals = []
                for col in df.columns:
                    v = row[col]
                    if v is None or (isinstance(v, float) and v != v):
                        vals.append("NULL")
                    elif isinstance(v, (int, float)):
                        vals.append(str(v))
                    else:
                        vals.append(f"'{str(v).replace(chr(39), chr(39)*2)}'")
                safe_cols = [c.replace(" ","_").replace("-","_") for c in df.columns]
                lines.append(f"INSERT INTO scraped_data ({', '.join(safe_cols)}) VALUES ({', '.join(vals)});")
            return "\n".join(lines).encode("utf-8"), "text/plain; charset=utf-8"

        raise ValueError(f"Unsupported buffer export format: {fmt}")

    def export_multiple(self, df: pd.DataFrame, formats: list[str] = None, prefix: str = "export") -> dict:
        if formats is None:
            formats = ["csv", "json", "excel"]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        results = {}
        for fmt in formats:
            fname = f"{prefix}_{timestamp}.{fmt}"
            try:
                if fmt == "csv":
                    results[fmt] = self.export_csv(df, fname)
                elif fmt == "excel":
                    results[fmt] = self.export_excel(df, fname)
                elif fmt == "json":
                    results[fmt] = self.export_json(df, fname)
                elif fmt == "word":
                    results[fmt] = self.export_word(df, fname)
                elif fmt == "pdf":
                    results[fmt] = self.export_pdf(df, f"{prefix}_{timestamp}.pdf")
                elif fmt == "html":
                    results[fmt] = self.export_html(df, fname)
                elif fmt == "parquet":
                    results[fmt] = self.export_parquet(df, fname)
                elif fmt == "xml":
                    results[fmt] = self.export_xml(df, fname)
                elif fmt == "sql":
                    results[fmt] = self.export_sql(df, "scraped_data", fname)
                else:
                    results[fmt] = {"error": f"Unknown format: {fmt}"}
            except Exception as e:
                results[fmt] = {"error": str(e)}
        return {
            "formats": formats,
            "results": results,
            "total_files": len([r for r in results.values() if "error" not in r]),
            "total_errors": len([r for r in results.values() if "error" in r]),
        }

    def get_export_stats(self, df: pd.DataFrame) -> dict:
        memory = df.memory_usage(deep=True).sum()
        return {
            "row_count": len(df),
            "column_count": len(df.columns),
            "memory_mb": round(memory / (1024 * 1024), 4),
            "columns": [
                {
                    "name": col,
                    "dtype": str(df[col].dtype),
                    "non_null": int(df[col].notna().sum()),
                    "null_pct": round(float(df[col].isna().mean() * 100), 2),
                    "unique": int(df[col].nunique()),
                }
                for col in df.columns
            ],
            "estimated_sizes": {
                "csv": self._estimate_size(df, "csv"),
                "json": self._estimate_size(df, "json"),
                "excel": self._estimate_size(df, "excel"),
                "word": self._estimate_size(df, "word"),
                "parquet": self._estimate_size(df, "parquet"),
            },
        }

    def _estimate_size(self, df: pd.DataFrame, fmt: str) -> str:
        if fmt == "csv":
            bytes_est = len(df.to_csv(index=False).encode())
        elif fmt == "json":
            bytes_est = len(json.dumps(df.to_dict(orient="records"), default=str).encode())
        elif fmt == "parquet":
            buffer = df.to_parquet(index=False)
            bytes_est = len(buffer) if buffer else 0
        else:
            bytes_est = df.memory_usage(deep=True).sum()
        if bytes_est < 1024:
            return f"{bytes_est} B"
        elif bytes_est < 1024 * 1024:
            return f"{bytes_est / 1024:.1f} KB"
        else:
            return f"{bytes_est / (1024 * 1024):.1f} MB"
